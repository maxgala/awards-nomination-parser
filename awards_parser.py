""" Read from a formatted Excel doc and create a Word doc with information based on award being nominated for
"""

from openpyxl import load_workbook
from openpyxl.formula import Tokenizer
from docx import Document

social_media = 0

def main():
	workbook = load_workbook(filename='Awards_2019_Parsing.xlsx')
	field_by_award_ws = workbook["Field by Award"]
	workbook = load_workbook(filename='2019 MAX Awards Nomination Form.xlsx')
	submissions_ws = workbook["Sheet1"]
	fields = field_by_award_ws["B"]
	global social_media
	
	for row in submissions_ws.iter_rows(min_row=2, max_row=submissions_ws.max_row, max_col=submissions_ws.max_column):
		award_type = row[1].value
		for column in field_by_award_ws.iter_cols(max_row=field_by_award_ws.max_row, max_col=field_by_award_ws.max_column, min_col=2):
			if column[0].value.lower() == award_type.lower():
				social_media = 0
				make_doc(row[0].row, column, award_type, fields, field_by_award_ws, submissions_ws)

def print_to_doc(cell, field, document, last_paragraph):
	data = None	
	if cell.value:
		data = str(cell.value)
	else:
		return last_paragraph
	if field == 'Name (First Name)':
		last_paragraph.add_run('Nominee Information').bold = True
		paragraph = document.add_paragraph('Full Name: ')
		paragraph.add_run(data)
		return paragraph
	elif field == 'Organization Contact Name (First Name)':
		paragraph = document.add_paragraph()
		paragraph.add_run('Organization Contact Name: ')
		paragraph.add_run(data)
		return 
	elif field == 'Name (Last Name)' or field == 'Organization Contact Name (Last Name)':
		last_paragraph.add_run(' ')
		last_paragraph.add_run(data)
		return 
	elif field == 'Email' or field == 'Phone Number':
		paragraph = document.add_paragraph()
		paragraph.add_run(field)
		paragraph.add_run(': ')
		paragraph.add_run(data)
		return 
	elif field == 'Name of Company':
		paragraph = document.add_paragraph()
		paragraph.add_run(field).bold = True
		paragraph.add_run(': ').bold = True
		paragraph.add_run(data)
		return 
	elif field == 'LinkedIn:' or field == 'Social Media Profile:' or field == 'Other Public Profiles (e.g. website):':
		global social_media
		if social_media == 0:
			document.add_paragraph()
			last_paragraph = document.add_paragraph().add_run('Public Profile (e.g. LinkedIn, website)').bold = True
			social_media = 1
		last_paragraph = document.add_paragraph()
		last_paragraph.add_run(field)
		last_paragraph.add_run(' ')
		if "=" in data:
			tok = Tokenizer(data)
			for t in tok.items:
				if t.type == "OPERAND" and t.value != '""':			
					last_paragraph.add_run(t.value[1:len(t.value) - 1])
					return last_paragraph			
		last_paragraph.add_run(data)
		return last_paragraph
	elif "upload" in field.lower():		
		tok = Tokenizer(data)
		for t in tok.items:
			if t.type == "OPERAND" and t.value != '""':
				document.add_paragraph()
				paragraph = document.add_paragraph().add_run(field).bold = True				
				paragraph = document.add_paragraph(t.value[1:len(t.value) - 1])
				return paragraph
		return last_paragraph
	else:
		document.add_paragraph()
		document.add_paragraph().add_run(field).bold = True
		paragraph = document.add_paragraph(data)
		return paragraph

def make_doc(row_num, column, award_type, fields, fields_by_award_ws, submissions_ws):
	document = Document()
	paragraph_format = document.styles['Normal'].paragraph_format
	paragraph_format.space_after = 0
	
	award_type = 'MAX ' + award_type
	document.add_heading(award_type, level=0)
	last_paragraph = document.add_paragraph()
	file_name = "Submissions/"
	nominator_name = ""
	
	for cell in column:
		column_to_access = (fields_by_award_ws['A'][cell.row - 1]).value
		if cell.value == 'x' and fields[cell.row - 1].value != 'Award Being Nominated':
			last_paragraph = print_to_doc(submissions_ws[column_to_access + str(row_num)], fields[cell.row - 1].value, document, last_paragraph)
		if fields[cell.row - 1].value == 'Name (First Name)' or fields[cell.row - 1].value == 'Name (Last Name)':
			file_name += submissions_ws[column_to_access + str(row_num)].value
		elif fields[cell.row - 1].value == 'First Name' or fields[cell.row - 1].value == 'Last Name':
			if submissions_ws[column_to_access + str(row_num)].value != None:
				nominator_name += submissions_ws[column_to_access + str(row_num)].value
	if nominator_name != "":
		file_name += "_by_"
		file_name += nominator_name
	file_name += ".docx"
	document.save(file_name)

if __name__ == '__main__':
	main();